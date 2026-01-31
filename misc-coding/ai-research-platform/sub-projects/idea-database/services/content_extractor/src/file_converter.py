#!/usr/bin/env python3
"""
File Converter Module
Handles conversion of different file types to markdown format
"""

import os
import io
import logging
import tempfile
from typing import Optional, Dict, Any
from pathlib import Path

# Document processing imports
import PyPDF2
from docx import Document
from PIL import Image
import pytesseract
import magic

# Text processing imports
import re

logger = logging.getLogger(__name__)

class FileConverter:
    """Converts various file types to markdown format"""
    
    def __init__(self):
        self.supported_types = {
            'application/pdf': self.convert_pdf,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self.convert_docx,
            'image/jpeg': self.convert_image,
            'image/png': self.convert_image,
            'image/gif': self.convert_image,
            'image/bmp': self.convert_image,
            'image/tiff': self.convert_image,
            # HTML and text files will be handled by pre-processor service
        }
    
    async def convert_to_markdown(self, file_content: bytes, filename: str, mime_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Convert file content to markdown
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            mime_type: MIME type (will be detected if not provided)
        
        Returns:
            Dict with markdown content, metadata, and processing info
        """
        try:
            # Detect MIME type if not provided
            if not mime_type:
                mime_type = magic.from_buffer(file_content, mime=True)
            
            logger.info(f"Converting file to markdown: {filename} (type: {mime_type})")
            
            # Check if this is a text-based file that should be handled by pre-processor
            if mime_type in ['text/html', 'text/plain']:
                return {
                    'success': True,
                    'requires_preprocessing': True,
                    'raw_content': file_content.decode('utf-8', errors='ignore'),
                    'metadata': {
                        'original_filename': filename,
                        'mime_type': mime_type,
                        'file_size': len(file_content)
                    }
                }
            
            # Check if we support this file type for binary processing
            converter_func = self.supported_types.get(mime_type)
            if not converter_func:
                return {
                    'success': False,
                    'error': f'Unsupported file type: {mime_type}',
                    'markdown_content': '',
                    'metadata': {'original_filename': filename, 'mime_type': mime_type}
                }
            
            # Convert the file
            result = await converter_func(file_content, filename)
            
            return {
                'success': True,
                'markdown_content': result.get('content', ''),
                'metadata': {
                    'original_filename': filename,
                    'mime_type': mime_type,
                    'file_size': len(file_content),
                    'processing_info': result.get('info', {}),
                    'extracted_text_length': len(result.get('content', ''))
                }
            }
            
        except Exception as e:
            logger.error(f"Error converting file {filename}: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'markdown_content': '',
                'metadata': {'original_filename': filename, 'mime_type': mime_type}
            }
    
    async def convert_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Convert PDF to markdown"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            
            pages_content = []
            total_pages = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        pages_content.append(f"## Page {page_num + 1}\n\n{text.strip()}")
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1} from {filename}: {str(e)}")
                    pages_content.append(f"## Page {page_num + 1}\n\n*[Error extracting page content]*")
            
            markdown_content = f"# {Path(filename).stem}\n\n" + "\n\n---\n\n".join(pages_content)
            
            return {
                'content': markdown_content,
                'info': {
                    'total_pages': total_pages,
                    'pages_processed': len(pages_content),
                    'extraction_method': 'PyPDF2'
                }
            }
            
        except Exception as e:
            logger.error(f"PDF conversion error for {filename}: {str(e)}")
            return {
                'content': f"# {Path(filename).stem}\n\n*Error processing PDF: {str(e)}*",
                'info': {'error': str(e)}
            }
    
    async def convert_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Convert Word document to markdown"""
        try:
            doc = Document(io.BytesIO(file_content))
            
            content_parts = [f"# {Path(filename).stem}\n"]
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Basic formatting detection
                    if paragraph.style.name.startswith('Heading'):
                        level = int(paragraph.style.name.split(' ')[-1]) if paragraph.style.name.split(' ')[-1].isdigit() else 2
                        content_parts.append(f"{'#' * (level + 1)} {text}")
                    else:
                        content_parts.append(text)
                    paragraph_count += 1
            
            # Process tables
            table_count = 0
            for table in doc.tables:
                table_md = "\n\n| "
                headers = []
                for cell in table.rows[0].cells:
                    headers.append(cell.text.strip() or "Column")
                table_md += " | ".join(headers) + " |\n|"
                table_md += " --- |" * len(headers) + "\n"
                
                for row in table.rows[1:]:
                    table_md += "| "
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip().replace('\n', ' '))
                    table_md += " | ".join(row_data) + " |\n"
                
                content_parts.append(table_md)
                table_count += 1
            
            markdown_content = "\n\n".join(content_parts)
            
            return {
                'content': markdown_content,
                'info': {
                    'paragraphs': paragraph_count,
                    'tables': table_count,
                    'extraction_method': 'python-docx'
                }
            }
            
        except Exception as e:
            logger.error(f"DOCX conversion error for {filename}: {str(e)}")
            return {
                'content': f"# {Path(filename).stem}\n\n*Error processing Word document: {str(e)}*",
                'info': {'error': str(e)}
            }
    
    async def convert_image(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Convert image to markdown using OCR"""
        try:
            # Save to temporary file for pytesseract
            with tempfile.NamedTemporaryFile(suffix=Path(filename).suffix, delete=False) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                # Open and preprocess image
                image = Image.open(temp_file_path)
                
                # Convert to RGB if necessary
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                
                # Extract text using OCR
                extracted_text = pytesseract.image_to_string(image)
                
                # Get confidence data
                ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
                confidences = [int(conf) for conf in ocr_data['conf'] if int(conf) > 0]
                avg_confidence = sum(confidences) / len(confidences) if confidences else 0
                
                # Format as markdown
                markdown_content = f"# {Path(filename).stem}\n\n"
                markdown_content += f"![{filename}]({filename})\n\n"
                
                if extracted_text.strip():
                    markdown_content += "## Extracted Text\n\n"
                    markdown_content += extracted_text.strip()
                else:
                    markdown_content += "*No text detected in image*"
                
                return {
                    'content': markdown_content,
                    'info': {
                        'image_size': f"{image.width}x{image.height}",
                        'ocr_confidence': round(avg_confidence, 2),
                        'text_length': len(extracted_text.strip()),
                        'extraction_method': 'pytesseract'
                    }
                }
                
            finally:
                # Clean up temporary file
                os.unlink(temp_file_path)
                
        except Exception as e:
            logger.error(f"Image conversion error for {filename}: {str(e)}")
            return {
                'content': f"# {Path(filename).stem}\n\n![{filename}]({filename})\n\n*Error processing image: {str(e)}*",
                'info': {'error': str(e)}
            } 